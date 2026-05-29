"""Generate comprehensive ARCHITEX-CAD project outline Word document."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx import Document
from docx_helpers import (
    set_doc_defaults,
    add_title,
    add_h1,
    add_h2,
    add_h3,
    add_para,
    add_bullets,
    add_table,
    add_footer,
)

OUT = Path(__file__).resolve().parent.parent / "docs" / "ARCHITEX-CAD_Comprehensive_Project_Outline.docx"


def build():
    doc = Document()
    set_doc_defaults(doc)

    add_title(
        doc,
        "ARCHITEX-CAD (InFra_TeCh)",
        "Comprehensive Project Outline — Technical Architecture & Capabilities",
    )

    # TABLE OF CONTENTS (manual list — Word TOC field not set via python-docx easily)
    add_h1(doc, "Table of Contents")
    toc = [
        "1. Executive Summary",
        "2. Project Identity & Purpose",
        "3. High-Level Architecture",
        "4. Technology Stack",
        "5. Repository Structure",
        "6. Application Flow & Workspace Model",
        "7. Frontend Architecture (React/TypeScript)",
        "8. Python Calculation Backend (FastAPI)",
        "9. Engineering Capabilities by Domain",
        "10. Calculator Modules (Frontend)",
        "11. API Route Reference (Backend)",
        "12. State Management (Zustand Stores)",
        "13. Services Layer",
        "14. BIM/CAD Viewer Capabilities",
        "15. Electron Desktop Shell",
        "16. Mobile Companion Application",
        "17. Shared Package & Geometry Extensions",
        "18. Regional & African Market Focus",
        "19. End-to-End Workflows",
        "20. Dependencies & Prerequisites",
        "21. Scope Boundaries & Limitations",
        "22. Appendices",
    ]
    add_numbered = lambda items: [doc.add_paragraph(i, style="List Number") for i in items]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1
    add_h1(doc, "1. Executive Summary")
    add_para(
        doc,
        "ARCHITEX-CAD (npm package name: infraafrica) is an integrated civil and structural "
        "engineering workstation designed for African infrastructure delivery. It unifies BIM "
        "visualization, parametric sketching, standards-based engineering calculations, site "
        "intelligence, bill of quantities (BoQ), regulatory documents, government portfolio "
        "management, and field data capture into a single desktop application backed by a Python "
        "FastAPI calculation engine with over 200 HTTP endpoints."
    )
    add_para(
        doc,
        "The platform targets engineers working on buildings, roads, water and sanitation (WASH), "
        "and energy projects across Sub-Saharan Africa, with particular localization for Zambia "
        "(ZMW rates, EIZ memos, black cotton soils, ZEMA/EIA context) while supporting international "
        "design standards including Eurocodes, BS 8110, AASHTO, EC8, and IEC electrical standards."
    )

    # 2
    add_h1(doc, "2. Project Identity & Purpose")
    add_h2(doc, "2.1 Product Names")
    add_table(
        doc,
        ["Name", "Usage"],
        [
            ("ARCHITEX-CAD", "Product / user-facing brand"),
            ("infraafrica", "npm package identifier"),
            ("InFra_TeCh", "Workspace / folder name"),
        ],
    )
    add_h2(doc, "2.2 Core Value Proposition")
    add_bullets(
        doc,
        [
            "Single integrated environment replacing fragmented spreadsheets, GIS, BoQ tools, and viewers.",
            "BIM-to-calculation-to-cost pipeline with auditable step-by-step design traces.",
            "Location-based site intelligence (terrain, soil, climate, seismic) from GPS coordinates.",
            "Africa-first localization: ZMW pricing, regional tenders, EIA screening, local material context.",
            "Field-to-office sync via mobile companion application.",
        ],
    )
    add_h2(doc, "2.3 Target Users")
    add_bullets(
        doc,
        [
            "Civil and structural engineers",
            "Geotechnical engineers",
            "Highway and pavement engineers",
            "Water and sanitation engineers",
            "Electrical engineers on multidisciplinary infrastructure projects",
            "Government agencies and consultants managing public infrastructure portfolios",
            "Quantity surveyors and project managers requiring BoQ and EVM tools",
        ],
    )

    # 3
    add_h1(doc, "3. High-Level Architecture")
    add_para(doc, "The system follows a three-tier local architecture:")
    add_table(
        doc,
        ["Tier", "Technology", "Responsibility"],
        [
            ("Presentation", "Electron + React 18 + TypeScript + Vite + Tailwind", "UI, BIM viewer, calculators, panels"),
            ("Application / API", "Python 3.11 + FastAPI + Uvicorn", "Engineering calculations, BIM parse, documents, geo APIs"),
            ("Field", "React Native (Expo)", "Site reports, photos, checklists, quick calcs, sync"),
        ],
    )
    add_h2(doc, "3.1 Data Flow")
    add_bullets(
        doc,
        [
            "Electron main process spawns Python server on localhost:8000 (unless external mode).",
            "React SPA calls FastAPI via HTTP (VITE_API_BASE, default http://127.0.0.1:8000).",
            "IFC models parsed client-side (web-ifc WASM) and/or server-side (ifcopenshell).",
            "Meshes loaded into xeokit 3D viewer; selection bridges to calculator inputs.",
            "Mobile pushes field data to POST /sync/receive when on same network as desktop.",
            "Offline: Electron SQLite cache for projects and calculations when server unavailable.",
        ],
    )
    add_h2(doc, "3.2 Calculation Architecture (Two Layers)")
    add_table(
        doc,
        ["Layer", "Path", "Role"],
        [
            ("Standards engines", "python/calculations/", "Eurocode/BS/AASHTO engines with step traces and clause refs"),
            ("API adapters", "python/calculators/", "Pydantic wrappers and chart/time-series simulations for UI"),
        ],
    )

    # 4
    add_h1(doc, "4. Technology Stack")
    add_h2(doc, "4.1 Desktop & Frontend")
    add_table(
        doc,
        ["Component", "Version / Library", "Purpose"],
        [
            ("Runtime", "Electron 28+", "Desktop shell, file dialogs, IPC, offline DB"),
            ("UI Framework", "React 18", "Single-page application"),
            ("Language", "TypeScript 5.3", "Type-safe frontend code"),
            ("Build", "Vite 8", "Dev server and production bundle"),
            ("Styling", "Tailwind CSS 3.4", "Utility-first CSS"),
            ("State", "Zustand 4.4", "~25 domain stores"),
            ("BIM Viewer", "@xeokit/xeokit-sdk 2.6", "3D model rendering"),
            ("IFC Parse", "web-ifc 0.0.68", "Client-side IFC WASM parsing"),
            ("3D Math", "three.js 0.184", "Sketch meshes, overlays"),
            ("GIS", "Leaflet 1.9", "2D map viewer"),
            ("Diagrams", "@xyflow/react", "SLD and flow diagrams"),
            ("i18n", "i18next + react-i18next", "English/French locales"),
            ("Local DB", "better-sqlite3", "Electron offline persistence"),
        ],
    )
    add_h2(doc, "4.2 Backend")
    add_table(
        doc,
        ["Component", "Purpose"],
        [
            ("FastAPI + Uvicorn", "HTTP API server"),
            ("NumPy + SciPy", "FEA, optimization, numerics"),
            ("Optuna", "Generative design optimization"),
            ("pandapower", "Power flow, short-circuit analysis"),
            ("ifcopenshell + ezdxf + shapely", "BIM, CAD, 2D geometry"),
            ("reportlab + Pillow", "PDF reports, image processing"),
            ("openpyxl", "BoQ Excel export"),
            ("anthropic", "AI chat (Claude)"),
            ("Optional: wntr", "EPANET pipe network hydraulics"),
            ("Optional: rasterio/pyproj", "Advanced GIS (GDAL stack)"),
            ("Optional: openseespy", "Extended FEA (non-Windows)"),
        ],
    )
    add_h2(doc, "4.3 Mobile")
    add_bullets(doc, ["React Native 0.85", "Expo SDK 56", "Offline SQLite", "GPS camera integration"])

    # 5
    add_h1(doc, "5. Repository Structure")
    add_table(
        doc,
        ["Path", "Contents"],
        [
            ("src/", "React UI: components, stores, services, hooks, cad/, locales/"),
            ("python/", "FastAPI main.py, calculations/, calculators/, geo/, government/, documents/"),
            ("electron/", "main.js, preload, menu, offline-sync.js"),
            ("mobile/", "Expo field app screens and services"),
            ("packages/shared/", "Cross-platform constants (countries, checklists, schedule types)"),
            ("GeometryExtensions/", "Vendored AutoCAD .NET geometry helpers (R20/R25) — reference only"),
            ("scripts/", "Python install/start, native rebuild, doc generators"),
            ("docs/", "Generated documentation (Word outlines)"),
        ],
    )

    # 6
    add_h1(doc, "6. Application Flow & Workspace Model")
    add_h2(doc, "6.1 Navigation (No React Router)")
    add_bullets(
        doc,
        [
            "View state: dashboard | workspace (App.tsx)",
            "Dashboard: ProjectDashboard — open IFC/CAD, recent projects",
            "Workspace: lazy AppShell + global VoiceAssistant",
            "Project open: useProject().openIFC() → load model → switch to workspace",
        ],
    )
    add_h2(doc, "6.2 Workspace Panels (21)")
    add_para(doc, "WorkspacePanel type defines sidebar panels:")
    add_bullets(
        doc,
        [
            "viewer, calculator, boq, schedule, optimizer, seismic, geo, vision, ai",
            "realestate, government, documents, carbon, wash, energy, intelligence",
            "emerging, site, verification, materials, project",
        ],
    )
    add_h2(doc, "6.3 Main Canvas Views")
    add_table(
        doc,
        ["View", "Component", "Purpose"],
        [
            ("bim", "BIMViewer + CadToolsPanel", "3D model and sketch authoring"),
            ("gis", "GISViewer (Leaflet)", "GeoJSON/SHP site data"),
            ("sld", "SLDViewer", "Single-line electrical diagram"),
        ],
    )
    add_h2(doc, "6.4 Supported File Types")
    add_table(
        doc,
        ["Extensions", "Destination"],
        [
            ("IFC, DWG, DXF, STEP, STL, OBJ, GLTF", "BIM viewer"),
            ("GeoJSON, SHP", "GIS viewer"),
        ],
    )

    # 7
    add_h1(doc, "7. Frontend Architecture (React/TypeScript)")
    add_h2(doc, "7.1 Entry Point")
    add_para(doc, "main.tsx → ErrorBoundary → App.tsx. Loads i18n and global CSS.")
    add_h2(doc, "7.2 Major Component Folders")
    add_table(
        doc,
        ["Folder", "UI / Purpose"],
        [
            ("Layout/", "AppShell, Sidebar, TopBar, StatusBar, LanguageSwitcher"),
            ("Dashboard/", "Project picker, recent activity"),
            ("BIMViewer/", "~30 components: viewer, ribbon, tree, layers, inspector, quantities"),
            ("Calculator/", "CalculatorPanel, 43 modules, ResultsDisplay, ReportExporter"),
            ("BoQ/", "Quantity compile and export"),
            ("Schedule/", "4D BIM-linked timeline"),
            ("Optimizer/", "Generative structural/solar design"),
            ("Seismic/", "EC8 analysis and spectrum charts"),
            ("GeoIntelligence/, Geo/", "Site maps, consolidation, piles, slopes, tunnel RMR"),
            ("GIS/", "Leaflet map viewer"),
            ("AIDesign/, AI/", "Design brief, BIMCanvas, VoiceAssistant"),
            ("RealEstate/", "Valuation, feasibility, land use, mortgage"),
            ("Government/", "Portfolio dashboard, project tracker"),
            ("Documents/", "Tender, calc reports, EIA, ESG"),
            ("WASH/, Energy/", "Domain hubs and simulation panels"),
            ("Intelligence/", "Digital twin, IoT, predictive maintenance, collaboration"),
            ("Sustainability/, Carbon/", "Carbon calc and verifier"),
            ("Emerging/", "Satellite, drone, thermal, blockchain, AR, CV safety"),
            ("VisionCapture/", "Vision AI capture"),
            ("Site/", "ZambiaSitePanel"),
            ("Verification/", "QuantityVerifier"),
            ("Materials/", "MaterialPricePanel"),
            ("Structural/, Road/, Circuit/, Power/, Wind/, Environment/, FEA/", "Domain charts"),
        ],
    )
    add_h2(doc, "7.3 CAD Subsystem (src/cad/)")
    add_bullets(
        doc,
        [
            "ConstraintSolver.ts — geometric/dimensional constraints",
            "DOF analysis, feature tree, conflict alerts (ConflictAlert.tsx)",
            "Integration with drawStore, sketchConstraintStore, featureTreeStore",
        ],
    )
    add_h2(doc, "7.4 Hooks")
    add_bullets(
        doc,
        [
            "useProject — file open, extension routing, IFC load",
            "useBIMViewer — selection and load handlers",
            "useViewerShortcuts — keyboard shortcuts",
            "useToolActions — ribbon action dispatch",
            "useCalculation — calculator run integration",
        ],
    )

    doc.add_page_break()

    # 8
    add_h1(doc, "8. Python Calculation Backend (FastAPI)")
    add_h2(doc, "8.1 Entry Points")
    add_bullets(
        doc,
        [
            "python/main.py — primary FastAPI app (~200+ routes, WebSocket collaboration)",
            "python/cad/occ_routes.py — mounted at /occ/* (OpenCASCADE sketch kernel)",
        ],
    )
    add_h2(doc, "8.2 calculations/ Subpackages")
    add_table(
        doc,
        ["Package", "Key Modules", "Standards / Purpose"],
        [
            ("structural/", "beam, slab, column, foundation, steel, timber, masonry, crack_width, winkler", "EC2, BS 8110, BS 5628, EC3"),
            ("geo/", "bearing_capacity, settlement, slope_stability, african_soils", "Terzaghi, Meyerhof, Bishop"),
            ("geotechnical/", "borehole, black_cotton_soil", "Theis, expansive soil treatment"),
            ("wash/", "water_demand, borehole, sewer, pipe_network, treatment_plant, water_hammer", "WASH hydraulics"),
            ("roads/", "flexible_pavement, hydrology, geometric_design, traffic_load, gravel_road", "AASHTO, SATCC context"),
            ("pressure/", "foundation_bearing, lateral_earth, boussinesq, bridge_*, pavement, pipe, tank", "Pressure distributions"),
            ("loads/", "load_combinations, wind_loads", "EC0, ACI 318, BS 8110, EN wind"),
            ("energy/", "solar_pv, battery_storage", "Renewables sizing"),
            ("power/", "short_circuit, protection, harmonics", "IEC 60909, IDMT, harmonics"),
            ("fea/", "solver_2d, modal_analysis", "2D frame stiffness method"),
            ("seismic/", "response_spectrum", "EC8 Type 1, SRSS/CQC"),
            ("wind/", "panel_method", "2D panel CFD for buildings"),
            ("circuit/", "spice_solver", "DC, AC sweep, transient"),
            ("sustainability/", "carbon", "Embodied carbon RICS/ICE"),
            ("materials/", "selector", "Material recommendations"),
            ("boq/", "zmw_rates", "Zambia unit rates"),
            ("site/", "zambia_site_data", "Wind, seismic, rainfall by lat/lon"),
            ("project/", "project_store", "SQLite projects, calcs, documents"),
            ("reporting/", "eiz_memo", "EIZ/Lusaka council PDF"),
            ("generative/", "optimizer", "scipy/Optuna structural + solar"),
            ("core/", "engineer_control, calculation_db", "Review layer, persistence"),
        ],
    )
    add_h2(doc, "8.3 calculators/ Modules (API Wrappers & Simulations)")
    add_table(
        doc,
        ["Module", "Computes"],
        [
            ("energy_bess, energy_microgrid, energy_transmission", "BESS, microgrid VD, sag/tension"),
            ("energy_hydro, energy_biogas, energy_wind_wake", "Hydro, biogas, wake effects"),
            ("energy_grid_fault, energy_power_flow", "Fault levels, AC load flow"),
            ("energy_simulations, energy_extended_simulations", "24h solar-battery, profiles, catenary, fault decay"),
            ("wash_* (8 modules)", "Tower, EPANET, DEWATS, WTP, stormwater, landfill, irrigation, sims"),
            ("geo_* (7 modules)", "Piles, slope, consolidation, improvement, tunnel, sims, RMR"),
            ("structural_simulations, structural_more_simulations", "BMD/SFD, slab, P-M, wind façade"),
            ("road_simulations, environment_simulations", "Pavement stress, ESAL, hydrograph, landfill gas"),
            ("market_pricing", "Live material pricing for BoQ"),
        ],
    )
    add_h2(doc, "8.4 Supporting Python Packages")
    add_table(
        doc,
        ["Package", "Purpose"],
        [
            ("geo/", "geo_intelligence, terrain, soil, climate, seismic, geocoder, site_budget, cache"),
            ("government/", "portfolio_database, dashboard_engine, payment_certificates, reporting"),
            ("documents/", "tender_generator, calculation_report, eia_screening, esg_report"),
            ("bim/, boq/, ai/, vision/, intelligence/", "IFC, quantities, AI design, image analysis, digital twin"),
        ],
    )

    # 9
    add_h1(doc, "9. Engineering Capabilities by Domain")
    domains = [
        (
            "9.1 Structural & FEA",
            [
                "EC2 beams, slabs, columns, pad/strip foundations; BS 8110; BS 5628 masonry; EC3 steel; timber",
                "Load takedown, bearings, crack width, Winkler elastic foundation",
                "2D frame FEA, modal analysis, EC8 seismic spectrum",
                "Simulations: BMD/SFD, foundation pressure bulb, slab moments, P-M interaction, wind façade",
                "Engineer control: per-step review with clause references and pass/fail",
            ],
        ),
        (
            "9.2 Geotechnical & Site Intelligence",
            [
                "Bearing (Terzaghi/Meyerhof), settlement, slope (Bishop), site classification",
                "Piles, consolidation, ground improvement, tunneling/RMR, black cotton soil with ZMW costs",
                "Borehole: Theis drawdown, pump selection; gravel road design",
                "Geo APIs: combined site analysis, terrain/soil/climate/seismic, geocode, budget, cache",
                "african_soils.py, zambia_site_data.py",
            ],
        ),
        (
            "9.3 Roads & Transport",
            [
                "AASHTO flexible pavement (structural number)",
                "Rational method drainage + Manning; geometric design; ESAL traffic",
                "Simulations: layer stresses, ESAL growth, storm hydrographs",
            ],
        ),
        (
            "9.4 WASH",
            [
                "Population water demand, borehole, sewer design, pipe networks",
                "Treatment plants, DEWATS, water hammer, elevated tanks, stormwater, landfill, irrigation",
                "Optional WNTR/EPANET hydraulics; daily tower level and pipe pressure simulations",
            ],
        ),
        (
            "9.5 Energy & Power",
            [
                "Solar PV, BESS, microgrid VD, transmission sag/tension, hydro, biogas, wind wake",
                "Pandapower load flow, IEC 60909 short-circuit, IDMT grading, harmonics",
                "24h solar-battery sim, catenary, voltage-drop profiles, fault current decay",
            ],
        ),
        (
            "9.6 Pressure, Loads & Wind",
            [
                "Foundation bearing profiles, Rankine earth pressure, Boussinesq, consolidation stress",
                "Bridge hydrostatic/dynamic, pavement/pipe/tank pressures, EN wind distribution",
                "EC0/ACI/BS load combinations; 2D panel-method wind CFD",
            ],
        ),
        (
            "9.7 Seismic & Sustainability",
            [
                "EC8 elastic/design spectra; modal SRSS/CQC; building seismic response simulation",
                "Embodied carbon (RICS/ICE); carbon credits; ESG from BoQ",
            ],
        ),
        (
            "9.8 Electrical / Circuit",
            ["SPICE-lite DC, AC sweep, transient (/circuit/*)"],
        ),
        (
            "9.9 BIM, BoQ, Vision, AI",
            [
                "IFC parse/export, CAD DXF, geometry booleans, plan takeoff, DWG export",
                "BoQ extract/compile, live pricing, ZMW rates, Excel/PDF, verification compare",
                "Vision: photo → structure → CAD/report; AI: chat, design variants, voice, text-to-BIM",
            ],
        ),
        (
            "9.10 Government, Documents, Real Estate",
            [
                "Portfolio DB, EVM dashboard, payment certificates, cashflow/S-curves",
                "Tenders (ZM/KE/NG/GH), calc reports, EIA screening, ESG reports",
                "Plot valuation, feasibility, land-use optimization, mortgage",
            ],
        ),
        (
            "9.11 Digital Twin & Platform",
            [
                "Asset registry, IoT ingest, predictive maintenance, WebSocket collaboration",
                "4D schedule from BIM; flood D8, thermal sim; emerging stubs (satellite, drone, AR, blockchain)",
            ],
        ),
        (
            "9.12 CAD Geometry Kernel",
            ["/occ/* — OpenCASCADE sketch: offset, fillet, extrude, boolean, STEP export (Shapely fallback)"],
        ),
    ]
    for title, items in domains:
        add_h2(doc, title)
        add_bullets(doc, items)

    doc.add_page_break()

    # 10
    add_h1(doc, "10. Calculator Modules (Frontend)")
    add_para(doc, "43 files in src/components/Calculator/modules/ — wired in CalculatorPanel.tsx:")
    add_h3(doc, "Architectural (17)")
    add_bullets(
        doc,
        [
            "LoadCombinations, BeamCalculator, SlabCalculator, ColumnCalculator, FoundationCalculator",
            "BearingCalculator, LoadCalculator, WindCalculator, SteelCalculator, TimberCalculator",
            "MasonryCalculator, FEACalculator, CrackWidthCalculator, WinklerCalculator, RoadCalculator",
            "MaterialSelector; PressurePanel (pressure sub-panel)",
        ],
    )
    add_h3(doc, "Electrical (9)")
    add_bullets(
        doc,
        [
            "EnergyCalculator, MicrogridCalculator, TransmissionCalculator, HydroCalculator",
            "BiogasCalculator, WindWakeCalculator, GridFaultCalculator, CircuitCalculator, WindCFDCalculator",
        ],
    )
    add_h3(doc, "WASH (10)")
    add_bullets(
        doc,
        [
            "WashCalculator, WaterTowerCalculator, PipeNetworkCalculator, DewatsCalculator, WTPCalculator",
            "StormwaterCalculator, LandfillCalculator, IrrigationCalculator, WaterHammerCalculator, TankCalculator",
        ],
    )
    add_h3(doc, "Geo (8)")
    add_bullets(
        doc,
        [
            "GeoCalculator, PilesCalculator, SlopeStabilityCalculator, ConsolidationCalculator",
            "GroundImprovementCalculator, TunnelingCalculator, SeismicCalculator, BlackCottonCalculator",
        ],
    )

    # 11
    add_h1(doc, "11. API Route Reference (Backend)")
    add_para(doc, "Primary route groups in python/main.py:")
    add_table(
        doc,
        ["Prefix", "Purpose"],
        [
            ("/health", "Server health check"),
            ("/calculate/*", "Structural, FEA, loads, foundations, roads, energy, wind, carbon"),
            ("/wash/*, /calculate/wash/*", "WASH demand, networks, treatment, hammer, EPANET"),
            ("/geo/*", "Geotech calcs + site intelligence (terrain, soil, climate, seismic)"),
            ("/pressure/*", "Soil/bridge/pavement/pipe/tank pressure distributions"),
            ("/roads/*, /calculate/road/*", "Pavement, drainage, geometry, traffic, gravel"),
            ("/power/*, /circuit/*, /wind/*", "Short-circuit, relays, harmonics, SPICE, panel CFD"),
            ("/fea/*, /seismic/*", "2D FEA, modal, EC8 spectrum"),
            ("/structural/*", "Crack width, Winkler, simulation charts"),
            ("/energy/*, /api/energy/*", "BESS, microgrid, transmission, hydro, biogas, wake, fault, power flow"),
            ("/bim/*, /geometry/*", "IFC/CAD parse, booleans, DWG, plan takeoff"),
            ("/boq/*, /verification/*", "Quantities, compile, ZMW rates, compare"),
            ("/vision/*", "Image analysis, CAD/report generation"),
            ("/ai/*, /generate/bim", "Chat, design, voice, proposals, text-to-BIM"),
            ("/government/*", "Portfolio, EVM, certificates, cashflow, reports"),
            ("/documents/*", "Tender, calc report, EIA, ESG"),
            ("/real-estate/*", "Valuation, feasibility, land use, mortgage"),
            ("/intelligence/*", "Digital twin, predictive maintenance"),
            ("/collaboration/*", "WebSocket rooms"),
            ("/project/*, /site/*, /geotechnical/*", "Project store, EIZ export, Zambia site, black cotton"),
            ("/simulate/*, /optimize/*", "Flood, thermal, seismic, generative optimization"),
            ("/occ/*", "OpenCASCADE CAD kernel"),
            ("/sync/receive", "Mobile field data sync"),
        ],
    )
    add_para(doc, "Sample documented endpoints (README):")
    add_table(
        doc,
        ["Method", "Endpoint", "Description"],
        [
            ("GET", "/health", "Health check"),
            ("POST", "/calculate/beam", "Beam design EC2"),
            ("POST", "/calculate/slab", "Slab design"),
            ("POST", "/calculate/column", "Column design"),
            ("POST", "/calculate/foundation", "Foundation design"),
            ("POST", "/calculate/loads", "Load combinations"),
            ("POST", "/calculate/road/pavement", "Pavement design"),
            ("POST", "/calculate/road/drainage", "Drainage design"),
        ],
    )

    # 12
    add_h1(doc, "12. State Management (Zustand Stores)")
    add_table(
        doc,
        ["Store", "Role"],
        [
            ("workspaceStore", "Active panel, tabs, main view bim/gis/sld"),
            ("viewerStore", "Model path, selection, layers, tools, snap, grid"),
            ("projectStore", "Current project, recent projects"),
            ("calculationStore", "Active module, inputs, results, IFC bridge"),
            ("ifcModelStore", "Parsed IFC stats/elements"),
            ("drawStore / sketchConstraintStore / featureTreeStore", "Sketch and parametric CAD"),
            ("boqStore / scheduleStore", "BoQ and 4D schedule"),
            ("geoStore / washStore / energyStore / seismicStore / carbonStore", "Domain panel state"),
            ("aiStore / realEstateStore / governmentStore / intelligenceStore", "AI, RE, gov, twin"),
            ("engineerReviewStore / offlineSyncStore / designCodeStore", "Review, offline, code selection"),
        ],
    )

    # 13
    add_h1(doc, "13. Services Layer (src/services/)")
    add_h3(doc, "13.1 API Clients")
    add_bullets(
        doc,
        [
            "calculationAPI — 40+ POST endpoints (structural, road, wash, energy, geo, etc.)",
            "boqAPI, geoAPI, aiAPI, realEstateAPI, governmentAPI, documentsAPI",
            "tier2API, tier3API, platformAPI, bimGeometryAPI, geometryExtensionsAPI, occAPI",
            "optimizerAPI, scheduleAPI, pressureAPI, visionAPI, emergingAPI",
            "apiConfig.ts — VITE_API_BASE default http://127.0.0.1:8000",
        ],
    )
    add_h3(doc, "13.2 Client-Side Engines")
    add_bullets(
        doc,
        [
            "ifcParser, ifcMeshXeokit — IFC to xeokit scene",
            "viewerControls — camera, sections, measurements, explode, xray",
            "drawEngine, drawInteraction, sketchGeometry, sketchCadOps — sketch authoring",
            "sketchToBoQ, ifcQuantities, ifcBoqService — quantity pipelines",
            "timeline4d — schedule-driven visibility",
            "geoOverlayEngine, minimapEngine, selectionBridge",
            "collaborationWS, offlineCache, exportService, fileService",
        ],
    )

    doc.add_page_break()

    # 14
    add_h1(doc, "14. BIM/CAD Viewer Capabilities")
    add_h2(doc, "14.1 Core Viewer")
    add_bullets(
        doc,
        [
            "@xeokit/xeokit-sdk with NavCube, TreeView, SectionPlanes, Measurements, MarqueePicker",
            "Formats: IFC, DWG/DXF (server), STEP/STL/OBJ/GLTF/FBX",
            "4D schedule visibility; CollaborationPresence + WebSocket",
        ],
    )
    add_h2(doc, "14.2 Navigation & Display")
    add_bullets(
        doc,
        [
            "Orbit, pan, zoom, walk; perspective/plan/ortho; fit/reset",
            "Explode, x-ray, isolate, highlight, storey cycle, sun study, snap, grid, minimap",
        ],
    )
    add_h2(doc, "14.3 Sketch & BIM Authoring")
    add_bullets(
        doc,
        [
            "Draw: line, polyline, wall, slab, column, rect, polygon, pipe, site boundary",
            "Extrude, move, rotate, copy, mirror, array, export IFC",
            "ConstraintPanel, DOFIndicator, ConflictAlert, FeatureTreePanel",
        ],
    )
    add_h2(doc, "14.4 AutoCAD-Style Tools")
    add_bullets(
        doc,
        [
            "2D: circle, arc, ellipse, hatch, spline, xline, region, donut, revcloud",
            "Modify: trim, extend, offset, fillet, chamfer, join, break, align, erase",
            "Layers, dimensions, mtext, blocks, xref",
            "3D: presspull, revolve, sweep, loft, union/subtract/intersect, slice",
        ],
    )
    add_h2(doc, "14.5 BIM Workflow Ribbon")
    add_bullets(
        doc,
        [
            "BoQ extract/import/export, 4D play/pause, quantities, plan takeoff",
            "Boolean diff, model compare, clash detection, DWG export",
            "ESG/carbon, tender, EIA, AI variants, calc report, marketplace BoQ",
        ],
    )
    add_h2(doc, "14.6 Geo Overlays in Viewer")
    add_bullets(
        doc,
        ["Terrain, soil, climate, seismic, contours, flood, polygon area, region cut/union, EIA"],
    )

    # 15
    add_h1(doc, "15. Electron Desktop Shell")
    add_bullets(
        doc,
        [
            "electron/main.js — BrowserWindow, spawns Python uvicorn on port 8000",
            "Native menu (menu.js), IPC preload bridge",
            "File open/save dialogs for IFC and project files",
            "offline-sync.js — SQLite: projects, calculations, sync queue",
            "Packaging: electron-builder — NSIS (Win), DMG (Mac), AppImage (Linux)",
            "Product ID: com.architex.cad; Product name: ARCHITEX-CAD",
        ],
    )

    # 16
    add_h1(doc, "16. Mobile Companion Application")
    add_table(
        doc,
        ["Screen / Feature", "Capability"],
        [
            ("SiteReportScreen", "Daily site reports, offline-first"),
            ("camera.ts", "GPS-tagged photo capture"),
            ("ChecklistScreen", "Foundation/structural phases (shared constants)"),
            ("CalculatorScreen", "Quick concrete, rebar, beam calcs"),
            ("SyncScreen", "POST to desktop /sync/receive on WiFi"),
            ("sqlite.ts", "Local offline storage"),
        ],
    )

    # 17
    add_h1(doc, "17. Shared Package & Geometry Extensions")
    add_h2(doc, "17.1 @infraafrica/shared")
    add_bullets(
        doc,
        [
            "APP_NAME, AFRICAN_COUNTRIES (10 countries: ZM, KE, NG, GH, TZ, ZW, BW, MZ, SN, CI)",
            "CHECKLIST_PHASES, FOUNDATION_CHECKLIST",
            "ScheduleActivity, ScheduleResult types",
        ],
    )
    add_h2(doc, "17.2 GeometryExtensions/")
    add_para(
        doc,
        "Vendored AutoCAD .NET geometry helper assemblies (R20 for AutoCAD 2015–2024, R25 for 2025+). "
        "Reference library for CAD interoperability — not part of the web application runtime."
    )

    # 18
    add_h1(doc, "18. Regional & African Market Focus")
    add_table(
        doc,
        ["Area", "Implementation"],
        [
            ("Zambia", "ZMW BoQ rates, black cotton module, EIZ memos, Zambia site panel, ZPPA export"),
            ("Multi-country", "Tender/EIA per ZM, KE, NG, GH; 10 countries in shared constants"),
            ("Soils", "african_soils.py, expansive soil costing in ZMW"),
            ("AI context", "Local manufacturers (Lafarge, Metlika, Kafue Brickworks, etc.)"),
            ("Standards", "SATCC roads, Kenya RDM, Nigerian codes, Zambia Building Regulations in AI prompts"),
        ],
    )

    # 19
    add_h1(doc, "19. End-to-End Workflows")
    workflows = [
        (
            "19.1 Design Office — Building",
            "Open IFC → EC2 beam/slab/column with engineer review → export calc PDF → "
            "compile BoQ in ZMW → tender pack + EIZ memo.",
        ),
        (
            "19.2 Site Selection",
            "Enter coordinates → geo intelligence → bearing/settlement → black cotton treatment costing → site budget.",
        ),
        (
            "19.3 Infrastructure — Roads",
            "Pavement AASHTO → drainage → ESAL → storm hydrograph → BoQ quantities.",
        ),
        (
            "19.4 WASH Project",
            "Population demand → borehole → pipe network (EPANET) → treatment → stormwater → BoQ + ESG.",
        ),
        (
            "19.5 Renewable Site",
            "Solar/BESS → transmission sag → grid fault → SLD in viewer → 24h simulation.",
        ),
        (
            "19.6 Government Programme",
            "Portfolio dashboard → EVM → payment certificates → standardized reports per typology.",
        ),
        (
            "19.7 Field Operations",
            "Mobile checklist + GPS photos → sync to desktop → attach to project record.",
        ),
    ]
    for title, text in workflows:
        add_h2(doc, title)
        add_para(doc, text)

    # 20
    add_h1(doc, "20. Dependencies & Prerequisites")
    add_h2(doc, "20.1 Development Prerequisites")
    add_bullets(
        doc,
        [
            "Node.js 18+ and npm 10+",
            "Python 3.11+ (Add to PATH on Windows)",
            "Optional: MSVC build tools for wntr on Windows",
            "Optional: OSGeo4W for rasterio/GDAL on Windows",
        ],
    )
    add_h2(doc, "20.2 Quick Start Commands")
    add_table(
        doc,
        ["Command", "Action"],
        [
            ("npm install", "Install Node dependencies"),
            ("py -3 -m pip install -r python/requirements.txt", "Install Python dependencies"),
            ("npm run python:dev", "Start calculation server"),
            ("npm run dev", "Start Vite frontend (browser)"),
            ("npm run electron:dev", "Full desktop app (Vite + Python + Electron)"),
            ("npm run electron:build", "Production desktop installer"),
        ],
    )
    add_h2(doc, "20.3 Requirements Files")
    add_bullets(
        doc,
        [
            "python/requirements.txt — core FastAPI stack",
            "python/requirements-geo.txt — optional GDAL/GIS",
            "python/requirements-wntr.txt — optional EPANET hydraulics",
        ],
    )

    # 21
    add_h1(doc, "21. Scope Boundaries & Limitations")
    add_bullets(
        doc,
        [
            "Workstation platform — not a certified substitute for independent checking on every project.",
            "Not full AutoCAD replacement — ribbon mimics AutoCAD; kernel is xeokit + server OCC.",
            "Not cloud SaaS by default — local Electron + localhost Python.",
            "Some features require optional installs (EPANET, OpenCASCADE, GDAL, openseespy).",
            "Emerging modules (satellite, blockchain, AR) are integration-ready stubs in places.",
            "Professional judgment and peer review remain essential for all outputs.",
        ],
    )

    # 22 Appendices
    add_h1(doc, "22. Appendices")
    add_h2(doc, "Appendix A — npm Scripts")
    add_table(
        doc,
        ["Script", "Description"],
        [
            ("dev", "Vite dev server"),
            ("build", "tsc && vite build"),
            ("electron:dev", "Concurrent Vite + Python + Electron"),
            ("electron:build", "Production installer"),
            ("python:dev", "Start Python server only"),
            ("python:status", "Check /health endpoint"),
        ],
    )
    add_h2(doc, "Appendix B — Verification Scripts (python/)")
    add_bullets(
        doc,
        [
            "verify_suite.py, verify_platform.py, verify_geometry.py",
            "verify_geometry_extensions.py, verify_tier23.py, verify_blueprint8.py",
            "verify_boq_geo.py, test_imports.py",
        ],
    )
    add_h2(doc, "Appendix C — Document Generation")
    add_para(
        doc,
        "This document was generated by scripts/generate_comprehensive_outline_docx.py. "
        "Civil engineer portfolio: scripts/generate_portfolio_docx.py → "
        "docs/ARCHITEX-CAD_Civil_Engineer_Portfolio.docx"
    )

    add_footer(
        doc,
        "ARCHITEX-CAD (InFra_TeCh) — Comprehensive Project Outline — Generated for technical documentation.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
