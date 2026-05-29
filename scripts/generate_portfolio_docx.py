"""Generate ARCHITEX-CAD civil engineer portfolio Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).resolve().parent.parent / "docs" / "ARCHITEX-CAD_Civil_Engineer_Portfolio.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_defaults(doc)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title(
        doc,
        "ARCHITEX-CAD",
        "Integrated Civil Engineering Platform — Portfolio Overview for Civil Engineers",
    )

    add_para(
        doc,
        "A desktop engineering workstation built for African infrastructure: design, analysis, "
        "BIM, costing, and project delivery in one place.",
        bold=False,
    )
    doc.add_paragraph()

    # --- 1 ---
    add_h1(doc, "1. The Problem It Addresses")
    add_para(
        doc,
        "On many infrastructure projects in Africa, engineers still work across disconnected tools: "
        "spreadsheets for structural checks, separate GIS for site data, another package for Bill of "
        "Quantities (BoQ), Word for reports, and a viewer for IFC models. Data does not flow between "
        "them, standards are applied inconsistently, and local context (soils, regulations, ZMW rates) "
        "is often added manually.",
    )
    add_para(
        doc,
        "ARCHITEX-CAD was built to bring that workflow into one integrated environment — so a civil "
        "engineer can move from site coordinates → geotechnical assumptions → structural design → "
        "quantities → tender documents without re-entering the same information five times.",
    )

    # --- 2 ---
    add_h1(doc, "2. What It Is, in Plain Terms")
    add_para(doc, "Think of it as a professional engineering desk on your computer:")
    add_bullets(
        doc,
        [
            "You open a 3D building or infrastructure model (IFC and other formats).",
            "You run code-based calculations (Eurocode 2, BS 8110, geotechnical, roads, water, power) "
            "and see step-by-step design checks, not just a final number.",
            "You extract quantities and costs (including Zambia-specific rates in ZMW).",
            "You produce reports, tenders, and compliance documents aligned with local practice.",
            "A mobile companion captures site photos, checklists, and quick field calculations, "
            "then syncs back to the desktop when WiFi is available.",
        ],
    )
    add_para(
        doc,
        "It is aimed especially at engineers working on buildings, roads, water/sanitation, and energy "
        "projects in Zambia and wider Sub-Saharan Africa, though many modules follow international "
        "standards (Eurocodes, AASHTO, EC8, IEC).",
    )

    # --- 3 ---
    add_h1(doc, "3. Why a Civil Engineer Would Care")
    add_table(
        doc,
        ["Traditional Workflow", "With ARCHITEX-CAD"],
        [
            ("Model in one tool, design in Excel", "Model and design in the same workspace"),
            ("Site data from memory or old reports", "Enter GPS → terrain, soil, climate, seismic context"),
            ("“Black box” spreadsheet results", "Full calculation trace with pass/fail per step"),
            ("BoQ rebuilt from drawings by hand", "Quantities pulled from BIM or sketch geometry"),
            ("Separate tools for roads, water, structure", "One platform for structural, geo, roads, WASH, energy"),
        ],
    )
    add_para(
        doc,
        "The value is not only speed — it is traceability, consistency, and auditability, which matter "
        "for councils, EIZ submissions, and government portfolios.",
    )

    # --- 4 ---
    add_h1(doc, "4. Core Engineering Capabilities")

    add_h2(doc, "4.1 Structural & Building Design")
    add_bullets(
        doc,
        [
            "Beams, slabs, columns, pad and strip foundations — Eurocode 2 (with BS 8110 where applicable)",
            "Masonry (BS 5628), steel (EC3), timber, bearing pads, crack width",
            "Load combinations (EC0, ACI 318, BS 8110), wind loads",
            "Pressure diagrams: earth pressure, Boussinesq, foundation bearing profiles",
            "2D frame FEA and modal analysis; seismic design to EC8",
            "Charts: bending/shear diagrams, P–M interaction, foundation pressure, façade wind",
        ],
    )
    add_para(
        doc,
        "Each calculation is presented as an engineer-readable trail: inputs → intermediate values → "
        "code clause → pass or fail.",
    )

    add_h2(doc, "4.2 Geotechnical & Site Investigation Support")
    add_bullets(
        doc,
        [
            "Bearing capacity, settlement, slope stability, pile capacity, consolidation",
            "Ground improvement, tunneling / RMR, site classification, African soil context",
            "Black cotton (expansive) soil — treatment options with indicative ZMW costs",
            "Borehole drawdown and pump sizing",
            "Site intelligence from coordinates: elevation, slope, soil, rainfall, seismic",
        ],
    )

    add_h2(doc, "4.3 Roads & Pavement")
    add_bullets(
        doc,
        [
            "Flexible pavement design (AASHTO structural number)",
            "Drainage (rational method, Manning), geometric design, ESAL traffic loading",
            "Gravel road design (Zambia-oriented practice)",
            "Simulations: layer stresses, ESAL growth, storm hydrographs",
        ],
    )

    add_h2(doc, "4.4 Water, Sanitation & Environmental (WASH)")
    add_bullets(
        doc,
        [
            "Population water demand, borehole yield and drawdown",
            "Sewer and pipe networks (optional EPANET-style hydraulic analysis)",
            "Treatment plants, DEWATS, stormwater ponds, elevated tanks",
            "Water hammer, irrigation, landfill design",
            "Daily storage and pressure profile simulations",
        ],
    )

    add_h2(doc, "4.5 Energy & Electrical")
    add_bullets(
        doc,
        [
            "Solar PV and battery storage (BESS) sizing",
            "Microgrid voltage drop, transmission sag/tension",
            "Hydro, biogas, wind farm wake effects",
            "Grid fault levels, relay grading, harmonics",
            "Single-line diagram (SLD) view; 24-hour generation/storage simulations",
        ],
    )

    # --- 5 ---
    add_h1(doc, "5. BIM, Quantities, and Commercial Delivery")
    add_h3(doc, "BIM Viewer and Authoring")
    add_bullets(
        doc,
        [
            "Open IFC, DWG, DXF, STEP, STL, and other formats",
            "Measure distances, areas, volumes; section planes; isolate and explode",
            "Sketch walls, slabs, columns, site boundaries; AutoCAD-style modify tools",
            "4D scheduling: construction weeks linked to model visibility",
            "Clash detection, model comparison, plan takeoff",
        ],
    )
    add_h3(doc, "Bill of Quantities (BoQ)")
    add_bullets(
        doc,
        [
            "Extract quantities from BIM model or sketched geometry",
            "Apply unit rates including Zambia ZMW pricing",
            "Export to Excel and PDF; link to carbon / ESG reporting",
        ],
    )

    # --- 6 ---
    add_h1(doc, "6. Documents, Compliance, and Public-Sector Delivery")
    add_table(
        doc,
        ["Document Type", "Purpose"],
        [
            ("Structural calculation reports", "Professional narrative for submissions"),
            ("EIZ / Lusaka calculation memos", "PDF format for local regulatory context"),
            ("Tender packages", "Country-specific eligibility (ZM, KE, NG, GH)"),
            ("EIA preliminary screening", "ZEMA, NEMA, and similar thresholds"),
            ("ESG / embodied carbon reports", "Generated from BoQ quantities"),
        ],
    )
    add_h3(doc, "Government & Portfolio Management")
    add_bullets(
        doc,
        [
            "Portfolio dashboard — status, contract value, delays",
            "Earned value (EVM), S-curves, cashflow projections",
            "Interim and final payment certificates",
            "Project snapshots and variation tracking",
        ],
    )

    # --- 7 ---
    add_h1(doc, "7. Site Work: Mobile Companion")
    add_bullets(
        doc,
        [
            "Daily site reports (offline-first)",
            "GPS-tagged photos",
            "Inspection checklists (foundation, structural, roofing, finishes, handover)",
            "Quick calculators (concrete, rebar, beam)",
            "Sync to desktop via WiFi when on site LAN",
        ],
    )

    # --- 8 ---
    add_h1(doc, "8. Regional Focus (Sub-Saharan Africa)")
    add_table(
        doc,
        ["Area", "How It Shows Up"],
        [
            ("Zambia", "ZMW BoQ rates, black cotton soil, EIZ memos, Zambia site panel"),
            ("Sub-Saharan Africa", "Multi-country tenders and EIA; SATCC roads; local code awareness"),
            ("Soils & climate", "African soils logic; rainfall, wind, seismic from coordinates"),
            ("Procurement", "Tender generators; government portfolio tools"),
        ],
    )

    # --- 9 ---
    add_h1(doc, "9. Example Project Stories")
    add_h3(doc, "9.1 Two-Storey Clinic in Lusaka")
    add_para(
        doc,
        "Load IFC model → EC2 slab and column checks with full trace → black cotton foundation "
        "treatment costing → compile BoQ in ZMW → export EIZ memo and tender section.",
    )
    add_h3(doc, "9.2 Rural Water Scheme")
    add_para(
        doc,
        "Population demand → borehole drawdown → pipe network hydraulics → elevated tank sizing → "
        "daily level simulation → BoQ and ESG summary.",
    )
    add_h3(doc, "9.3 District Road Upgrade")
    add_para(
        doc,
        "Pavement design (AASHTO) → drainage → ESAL traffic → storm hydrograph → quantities for "
        "gravel and pavement layers.",
    )
    add_h3(doc, "9.4 Government Housing Programme")
    add_para(
        doc,
        "Multiple projects on portfolio dashboard → EVM and payment certificates → standardized calc "
        "reports per block type.",
    )

    # --- 10 ---
    add_h1(doc, "10. Strengths to Highlight in a Portfolio Pitch")
    add_bullets(
        doc,
        [
            "Breadth — structure, geo, roads, water, energy, BoQ, documents, government tracking",
            "Auditability — step-by-step calculations with code references",
            "Africa-first — ZMW, ZEMA/EIZ context, black cotton, local materials",
            "BIM-to-cost pipeline — model → quantities → rates → tender/ESG",
            "Field-to-office — mobile sync for site reality feeding the same project",
            "Scalable architecture — modular Python engines with rich desktop UI",
        ],
    )

    # --- 11 ---
    add_h1(doc, "11. Honest Scope Notes")
    add_bullets(
        doc,
        [
            "It is a workstation platform, not a certified substitute for independent checking on every job.",
            "Some advanced features (full EPANET, full OCC kernel) depend on optional installs.",
            "Primary strength is engineering calculation + BIM integration, not full production drafting replacement.",
        ],
    )

    # --- 12 ---
    add_h1(doc, "12. Elevator Pitch")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(
        "ARCHITEX-CAD is an integrated civil engineering platform for African infrastructure: open BIM "
        "models, run Eurocode and local geotechnical/road/water designs with full calculation trails, "
        "generate BoQs in ZMW, produce tender and EIZ-ready documents, and manage government portfolios "
        "— with a mobile app for site capture. It was built to reduce fragmentation between design, "
        "analysis, costing, and delivery, and to embed Sub-Saharan site context from day one."
    )
    run.italic = True

    doc.add_paragraph()
    add_h1(doc, "13. Technical Foundation (Brief)")
    add_bullets(
        doc,
        [
            "Desktop application (Electron) with React user interface",
            "Python calculation server (FastAPI) — 200+ engineering endpoints",
            "BIM engine: IFC libraries (web-ifc, xeokit)",
            "Standards libraries in Python with auditable step output",
            "Optional: OpenCASCADE geometry, pandapower for electrical, EPANET for networks",
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("ARCHITEX-CAD (InFra_TeCh) — Document generated for civil engineering portfolio use.")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fr.italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
